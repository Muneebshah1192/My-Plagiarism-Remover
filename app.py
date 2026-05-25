import os
import sqlite3
import io
import json
import random
import urllib.request
import urllib.error
from datetime import datetime, timedelta, date
from functools import wraps
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file, flash, Response
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

from text_engine import TOOL_CATEGORIES, TOOL_INFO, process_tool, metrics, jaccard_similarity, phrase_overlap, grammar_polish

APP_NAME = os.getenv('APP_NAME', 'TextForge Studio')
DB_PATH = os.getenv('DATABASE_PATH', os.path.join(os.path.dirname(__file__), 'instance', 'textforge.db'))
UPLOAD_EXTENSIONS = {'.txt', '.md', '.docx', '.pdf'}
PROJECT_CREATED_AT = os.getenv('PROJECT_CREATED_AT', 'May 25, 2026 at 01:18 AM PKT')
OWNER_NAME = os.getenv('OWNER_NAME', 'Syed Muneeb')
OWNER_EMAIL = os.getenv('OWNER_EMAIL', 'muneebshah1192@gmail.com')
OWNER_PHONE = os.getenv('OWNER_PHONE', '03140895219')
ADMIN_EMAIL = os.getenv('ADMIN_EMAIL', OWNER_EMAIL)
ADMIN_PASSWORD = os.getenv('ADMIN_PASSWORD', 'Muneebshah1192@theadmin.com')
PAYMENT_UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'instance', 'payment_proofs')

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'change-this-secret-key-before-production')
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

LEGAL_UPDATED = PROJECT_CREATED_AT
GUIDES = {
    'ai-humanizer-guide': {
        'title': 'AI Humanizer Guide: How to Make Text Clear, Natural, and Useful',
        'description': 'A practical guide explaining how rewriting, clarity, tone, and originality support tools help improve text quality.',
        'sections': [
            ('What an AI humanizer does', 'An AI humanizer helps transform stiff, repetitive, or robotic writing into clearer and more natural language. It should improve flow, reduce repeated phrases, and keep the original meaning understandable.'),
            ('How to use it responsibly', 'Use rewriting tools to polish your own drafts, improve readability, and create better structure. If ideas come from another source, add citations and avoid presenting copied research as your own.'),
            ('Best workflow', 'Paste your draft, choose a tool, review the output, compare similarity, then manually add examples, facts, and your own voice before publishing.')
        ]
    },
    'plagiarism-reduction-guide': {
        'title': 'Originality Support Guide: Reducing Text Overlap the Right Way',
        'description': 'Learn how text overlap is reduced through rewriting, sentence restructuring, citations, and proper editing.',
        'sections': [
            ('What originality support means', 'Originality support is not a promise to hide copied work. It means improving wording, structure, clarity, and presentation while encouraging proper attribution for borrowed ideas.'),
            ('Why rewriting alone is not enough', 'Changing words without understanding the idea can create weak writing. The strongest result comes from rewriting, adding examples, explaining in your own words, and citing sources.'),
            ('Practical checklist', 'Rewrite the paragraph, change sentence structure, remove repeated phrases, add a personal explanation, and include source references where needed.')
        ]
    },
    'seo-writing-tools-guide': {
        'title': 'SEO Writing Tools: Titles, Meta Descriptions, FAQs, and Keywords',
        'description': 'A helpful overview of text tools that support search-friendly writing and content planning.',
        'sections': [
            ('SEO content basics', 'Search-friendly writing needs a clear topic, useful headings, direct answers, readable formatting, and helpful supporting information.'),
            ('Useful SEO tools', 'Meta title generators, keyword extractors, FAQ builders, slug generators, and heading optimizers help writers structure pages more clearly.'),
            ('Content quality matters', 'SEO tools should support useful content, not replace it. Add accurate details, examples, and original insights to make the page valuable for readers.')
        ]
    },
    'prompt-generator-guide': {
        'title': 'Prompt Generator Guide: How to Create Better AI Prompts',
        'description': 'A simple guide for creating structured prompts for writing, video, marketing, study, and automation workflows.',
        'sections': [
            ('What makes a strong prompt', 'A strong prompt includes the goal, audience, style, format, constraints, examples, and expected output structure.'),
            ('Prompt structure', 'Start with the role, describe the task, add details, specify tone, set output format, and include quality rules.'),
            ('Common use cases', 'Prompt generators are useful for YouTube scripts, product descriptions, social captions, study notes, email drafts, and SEO content briefs.')
        ]
    }
}

CATEGORY_DESCRIPTIONS = {
    'Originality & AI Tools': 'Humanize, rewrite, clean AI-style writing, generate prompts, and reduce visible wording overlap responsibly.',
    'Smart Writing': 'Rewrite, humanize, simplify, expand, polish, and improve the natural flow of your text.',
    'Content Intelligence': 'Analyze writing quality, repetition, tone, readability, keywords, and similarity signals.',
    'SEO & Blogging': 'Create search-friendly titles, briefs, outlines, FAQs, metadata, headings, snippets, and content plans.',
    'Student Tools': 'Create notes, MCQs, flashcards, definitions, essay outlines, assignment formats, and study summaries.',
    'Business Tools': 'Draft emails, proposals, meeting notes, customer replies, resumes, cover letters, and business documents.',
    'Social Media': 'Generate hooks, captions, YouTube scripts, short video ideas, CTAs, hashtags, and social content calendars.',
    'Document Tools': 'Clean formatting, convert to Markdown, build outlines, extract action items, and export to PDF/DOCX.',
    'Premium Pro Tools': 'Advanced local tools for higher-value professional writing workflows. Premium or trial access is required.'
}

PREMIUM_TOOLS = {
    'plagiarism_remover','ai_humanizer','ai_detector_cleaner','deep_rewriter','originality_risk_report',
    'content_brief','blog_post_builder','seo_brief','youtube_script','short_video_script',
    'proposal_writer','resume_bullets','cover_letter','content_repurposer','executive_summary',
    'brand_voice_builder','landing_page_copy','sales_page_builder','ad_headline_lab','hook_strength_analyzer',
    'cta_optimizer','readability_fixer','content_scorecard','competitor_angle_generator','newsletter_writer'
}

# Premium aliases use existing algorithms so every listed premium tool works without an external API.
EXTRA_TOOLS = {
    'Premium Pro Tools': ['brand_voice_builder','landing_page_copy','sales_page_builder','ad_headline_lab','hook_strength_analyzer','cta_optimizer','readability_fixer','content_scorecard','competitor_angle_generator','newsletter_writer']
}
for _cat, _ids in EXTRA_TOOLS.items():
    TOOL_CATEGORIES.setdefault(_cat, [])
    for _id in _ids:
        if _id not in TOOL_CATEGORIES[_cat]:
            TOOL_CATEGORIES[_cat].append(_id)
TOOL_INFO.update({
    'brand_voice_builder': ('Brand Voice Builder', 'Create a brand voice guide with tone, audience, style rules, and sample phrases.'),
    'landing_page_copy': ('Landing Page Copy Builder', 'Create hero text, benefits, feature blocks, proof points, and CTA copy.'),
    'sales_page_builder': ('Sales Page Builder', 'Create persuasive sales-page copy with offer, benefits, objections, and CTA sections.'),
    'ad_headline_lab': ('Ad Headline Lab', 'Generate multiple direct-response ad headlines and angles.'),
    'hook_strength_analyzer': ('Hook Strength Analyzer', 'Score and improve hooks for social posts, videos, and ads.'),
    'cta_optimizer': ('CTA Optimizer', 'Turn weak calls-to-action into clearer, stronger conversion-focused CTAs.'),
    'readability_fixer': ('Readability Fixer', 'Rewrite text into a clearer, easier-to-read version.'),
    'content_scorecard': ('Content Scorecard', 'Create a practical scorecard for clarity, trust, originality, SEO, and actionability.'),
    'competitor_angle_generator': ('Content Angle Generator', 'Generate unique content angles, positioning ideas, and differentiation points.'),
    'newsletter_writer': ('Newsletter Writer', 'Turn a topic or notes into a structured newsletter draft.'),
})


def get_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def column_exists(db, table, column):
    cols = [row['name'] for row in db.execute(f'PRAGMA table_info({table})').fetchall()]
    return column in cols


def init_db():
    with get_db() as db:
        db.executescript('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            tool_id TEXT NOT NULL,
            input_text TEXT NOT NULL,
            output_text TEXT NOT NULL,
            stats_json TEXT NOT NULL,
            created_at TEXT NOT NULL,
            engine_used TEXT DEFAULT 'algorithm',
            FOREIGN KEY(user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS settings (
            user_id INTEGER PRIMARY KEY,
            default_tone TEXT DEFAULT 'professional',
            default_strength TEXT DEFAULT 'strong',
            brand_name TEXT DEFAULT 'TextForge Studio',
            engine_mode TEXT DEFAULT 'algorithm',
            gemini_api_key TEXT DEFAULT '',
            gemini_model TEXT DEFAULT 'gemini-1.5-flash',
            FOREIGN KEY(user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS site_settings (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS usage_logs (
            user_id INTEGER NOT NULL,
            usage_date TEXT NOT NULL,
            chars_used INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY(user_id, usage_date)
        );
        CREATE TABLE IF NOT EXISTS payment_proofs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            method TEXT NOT NULL,
            amount TEXT NOT NULL,
            currency TEXT DEFAULT 'USD',
            transaction_ref TEXT DEFAULT '',
            screenshot_path TEXT DEFAULT '',
            note TEXT DEFAULT '',
            status TEXT DEFAULT 'pending',
            created_at TEXT NOT NULL,
            reviewed_at TEXT DEFAULT '',
            admin_note TEXT DEFAULT '',
            FOREIGN KEY(user_id) REFERENCES users(id)
        );
        ''')
        user_columns = {
            'role': "TEXT NOT NULL DEFAULT 'user'",
            'plan': "TEXT NOT NULL DEFAULT 'free'",
            'trial_started_at': "TEXT DEFAULT ''",
            'premium_until': "TEXT DEFAULT ''",
        }
        for col, spec in user_columns.items():
            if not column_exists(db, 'users', col):
                db.execute(f"ALTER TABLE users ADD COLUMN {col} {spec}")
        setting_columns = {
            'engine_mode': "TEXT DEFAULT 'algorithm'",
            'gemini_api_key': "TEXT DEFAULT ''",
            'gemini_model': "TEXT DEFAULT 'gemini-1.5-flash'",
        }
        for col, spec in setting_columns.items():
            if not column_exists(db, 'settings', col):
                db.execute(f"ALTER TABLE settings ADD COLUMN {col} {spec}")
        defaults = {
            'ads_enabled': '0',
            'adsense_head_script': '',
            'ad_header_html': '',
            'ad_sidebar_html': '',
            'ad_footer_html': '',
            'ad_tool_html': '',
            'ads_txt_publisher_id': '',
            'site_domain': '',
            'plan_price_usd': '5',
            'trial_days': '1',
            'free_daily_chars': '4000',
            'trial_daily_chars': '20000',
            'premium_daily_chars': '100000',
            'premium_duration_days': '30',
            'enable_stripe': '0',
            'stripe_checkout_url': '',
            'enable_lemon': '0',
            'lemon_checkout_url': '',
            'enable_paddle': '0',
            'paddle_checkout_url': '',
            'enable_paypal': '0',
            'paypal_details': '',
            'enable_payoneer': '0',
            'payoneer_details': '',
            'enable_bank': '0',
            'bank_details': '',
            'enable_easypaisa': '1',
            'easypaisa_number': '0340-545-3639',
            'easypaisa_name': 'Manza Zahoor',
            'enable_sadapay': '1',
            'sadapay_number': '03140895219',
            'sadapay_name': 'syed Muneeb Haider Shah',
            'enable_mashreq': '0',
            'mashreq_details': '',
            'enable_allied': '0',
            'allied_details': '',
            'enable_api_mode': '1',
            'admin_notice': 'New users get a 1-day trial. Premium unlocks higher limits and Gemini API mode.',
        }
        for key, value in defaults.items():
            db.execute('INSERT OR IGNORE INTO site_settings(key,value) VALUES(?,?)', (key, value))
        # Keep existing sites compatible while correcting old default EasyPaisa details.
        old_easy = db.execute('SELECT value FROM site_settings WHERE key=?', ('easypaisa_number',)).fetchone()
        if old_easy and old_easy['value'] in ('0314-0340-545-3639', '031403405453639'):
            db.execute('UPDATE site_settings SET value=? WHERE key=?', ('0340-545-3639', 'easypaisa_number'))
        old_easy_name = db.execute('SELECT value FROM site_settings WHERE key=?', ('easypaisa_name',)).fetchone()
        if old_easy_name and old_easy_name['value'] in (OWNER_NAME, 'Syed Muneeb'):
            db.execute('UPDATE site_settings SET value=? WHERE key=?', ('Manza Zahoor', 'easypaisa_name'))
        admin = db.execute('SELECT * FROM users WHERE email=?', (ADMIN_EMAIL.lower(),)).fetchone()
        now = datetime.utcnow().isoformat()
        if not admin:
            cur = db.execute('INSERT INTO users(name,email,password_hash,created_at,role,plan,trial_started_at) VALUES(?,?,?,?,?,?,?)',
                (OWNER_NAME, ADMIN_EMAIL.lower(), generate_password_hash(ADMIN_PASSWORD), now, 'admin', 'premium', now))
            db.execute('INSERT OR IGNORE INTO settings(user_id, brand_name, engine_mode) VALUES(?,?,?)', (cur.lastrowid, APP_NAME, 'algorithm'))
        else:
            db.execute('UPDATE users SET role=?, plan=? WHERE email=?', ('admin', 'premium', ADMIN_EMAIL.lower()))
            db.execute('INSERT OR IGNORE INTO settings(user_id, brand_name, engine_mode) VALUES(?,?,?)', (admin['id'], APP_NAME, 'algorithm'))


@app.before_request
def ensure_db():
    init_db()


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return view(*args, **kwargs)
    return wrapped


def admin_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        user = current_user()
        if not user or user['role'] != 'admin':
            flash('Admin access is required.', 'error')
            return redirect(url_for('login'))
        return view(*args, **kwargs)
    return wrapped


def current_user():
    if 'user_id' not in session:
        return None
    with get_db() as db:
        return db.execute('SELECT * FROM users WHERE id=?', (session['user_id'],)).fetchone()


def site_settings():
    with get_db() as db:
        rows = db.execute('SELECT key,value FROM site_settings').fetchall()
    return {row['key']: row['value'] for row in rows}


def get_user_settings(user_id):
    with get_db() as db:
        row = db.execute('SELECT * FROM settings WHERE user_id=?', (user_id,)).fetchone()
        if not row:
            db.execute('INSERT INTO settings(user_id) VALUES(?)', (user_id,))
            row = db.execute('SELECT * FROM settings WHERE user_id=?', (user_id,)).fetchone()
    return row


def cat_slug(name):
    return name.lower().replace('&', 'and').replace(' ', '-').replace('/', '-').replace('--','-')


def category_by_slug(slug):
    for cat in TOOL_CATEGORIES:
        if cat_slug(cat) == slug:
            return cat
    return None


def build_tools_payload():
    return {
        cat: [
            {
                'id': tid,
                'name': TOOL_INFO.get(tid, (tid, ''))[0],
                'description': TOOL_INFO.get(tid, ('', ''))[1],
                'premium': tid in PREMIUM_TOOLS,
                'category': cat,
                'category_slug': cat_slug(cat),
            }
            for tid in ids if tid in TOOL_INFO
        ]
        for cat, ids in TOOL_CATEGORIES.items()
    }


def find_tool(tool_id):
    for cat, ids in TOOL_CATEGORIES.items():
        if tool_id in ids and tool_id in TOOL_INFO:
            return {
                'id': tool_id,
                'name': TOOL_INFO[tool_id][0],
                'description': TOOL_INFO[tool_id][1],
                'premium': tool_id in PREMIUM_TOOLS,
                'category': cat,
                'category_slug': cat_slug(cat),
            }
    return None


def int_setting(settings, key, default):
    try:
        return int(float(settings.get(key, default)))
    except Exception:
        return default


def parse_dt(value):
    if not value:
        return None
    try:
        return datetime.fromisoformat(value)
    except Exception:
        return None


def plan_status(user, settings=None):
    settings = settings or site_settings()
    now = datetime.utcnow()
    if not user:
        return {'name': 'guest', 'active': False, 'is_premium': False, 'is_trial': False, 'limit': int_setting(settings, 'free_daily_chars', 4000), 'days_left': 0}
    if user['role'] == 'admin':
        return {'name': 'admin', 'active': True, 'is_premium': True, 'is_trial': False, 'limit': 999999999, 'days_left': 3650}
    premium_until = parse_dt(user['premium_until'])
    if user['plan'] == 'premium' and premium_until and premium_until > now:
        days_left = max(0, (premium_until - now).days + 1)
        return {'name': 'premium', 'active': True, 'is_premium': True, 'is_trial': False, 'limit': int_setting(settings, 'premium_daily_chars', 100000), 'days_left': days_left}
    trial_started = parse_dt(user['trial_started_at'] or user['created_at'])
    trial_days = int_setting(settings, 'trial_days', 1)
    if trial_started and now <= trial_started + timedelta(days=trial_days):
        days_left = max(0, ((trial_started + timedelta(days=trial_days)) - now).days + 1)
        return {'name': 'trial', 'active': True, 'is_premium': False, 'is_trial': True, 'limit': int_setting(settings, 'trial_daily_chars', 20000), 'days_left': days_left}
    return {'name': 'free', 'active': False, 'is_premium': False, 'is_trial': False, 'limit': int_setting(settings, 'free_daily_chars', 4000), 'days_left': 0}


def today_usage(user_id):
    today = date.today().isoformat()
    with get_db() as db:
        row = db.execute('SELECT chars_used FROM usage_logs WHERE user_id=? AND usage_date=?', (user_id, today)).fetchone()
    return row['chars_used'] if row else 0


def add_usage(user_id, chars):
    today = date.today().isoformat()
    with get_db() as db:
        db.execute('INSERT INTO usage_logs(user_id, usage_date, chars_used) VALUES(?,?,?) ON CONFLICT(user_id, usage_date) DO UPDATE SET chars_used=chars_used+excluded.chars_used', (user_id, today, int(chars)))


def make_captcha():
    a = random.randint(2, 9)
    b = random.randint(2, 9)
    session['captcha_answer'] = str(a + b)
    return f'{a} + {b}'


def check_captcha():
    ans = request.form.get('captcha', '').strip()
    expected = session.get('captcha_answer', '')
    return expected and ans == expected


def fallback_tool_output(tool_id, text, params):
    # Route extra premium tools through existing strong local algorithms.
    tone = params.get('tone', 'professional') if params else 'professional'
    if tool_id == 'brand_voice_builder':
        return process_tool('content_brief', text, params) + '\n\nBrand Voice Guide:\n- Voice: clear, useful, trustworthy, and practical.\n- Tone: professional but friendly.\n- Avoid: exaggerated claims, unclear jargon, and robotic repetition.\n- Sample phrase: We make complex writing tasks easier, faster, and cleaner.'
    if tool_id == 'landing_page_copy':
        return process_tool('content_brief', text, params) + '\n\nLanding Page Copy:\n# Clearer writing. Faster workflows. Better results.\nTurn rough ideas into polished content with smart rewriting, SEO helpers, business tools, and document exports.\n\nBenefits:\n- Save time on everyday writing\n- Improve clarity and structure\n- Create copy for study, business, SEO, and social media\n\nCTA: Start creating better content today.'
    if tool_id == 'sales_page_builder':
        return process_tool('proposal_writer', text, params) + '\n\nSales Page Sections:\n1. Problem: writing takes time and quality matters.\n2. Solution: an all-in-one text workspace.\n3. Benefits: faster drafts, clearer communication, stronger structure.\n4. Offer: premium access with higher limits and advanced tools.\n5. CTA: Upgrade now.'
    if tool_id == 'ad_headline_lab':
        return process_tool('viral_hooks', text, params) + '\n\nAd Headlines:\n1. Write Better Content in Less Time\n2. Turn Rough Text Into Professional Copy\n3. Your All-in-One Writing Workspace\n4. Create SEO, Business, and Social Copy Faster\n5. Upgrade Your Writing Workflow Today'
    if tool_id == 'hook_strength_analyzer':
        base = process_tool('quality_score', text, params)
        return base + '\n\nHook Suggestions:\n- Start with a clear problem.\n- Add a benefit in the first line.\n- Use specific wording instead of generic claims.\n- Keep the opening under 15 words when possible.'
    if tool_id == 'cta_optimizer':
        return 'CTA Options:\n1. Start improving your writing today.\n2. Generate a polished version now.\n3. Upgrade to unlock premium writing workflows.\n4. Turn your text into professional content.\n5. Create your next draft in seconds.\n\nImproved CTA:\nStart your professional writing workflow now.'
    if tool_id == 'readability_fixer':
        return process_tool('simple_rewriter', text, {'tone':'simple', 'strength':'maximum'})
    if tool_id == 'content_scorecard':
        return process_tool('quality_score', text, params) + '\n\nScorecard:\n- Clarity: review sentence length and directness.\n- Trust: add proof, facts, or examples.\n- Originality: restructure long copied phrases.\n- SEO: add focused headings and keywords.\n- Actionability: include a clear next step.'
    if tool_id == 'competitor_angle_generator':
        keys = process_tool('keyword_extractor', text, params)
        return keys + '\n\nUnique Angles:\n1. Focus on speed and simplicity.\n2. Focus on professional output quality.\n3. Focus on student and freelancer use cases.\n4. Focus on no-API local tools.\n5. Focus on exports, history, and workflow organization.'
    if tool_id == 'newsletter_writer':
        return 'Subject: Useful writing tools to improve your workflow\n\nHi there,\n\nHere is a quick update built around your topic:\n\n' + process_tool('executive_summary', text, params) + '\n\nKey takeaways:\n- Improve clarity before publishing.\n- Use structured tools for faster drafting.\n- Review every output for accuracy and tone.\n\nThanks for reading,\nTextForge Studio'
    return process_tool(tool_id, text, params)


def make_gemini_prompt(tool_id, text, params):
    tool_name = TOOL_INFO.get(tool_id, (tool_id, ''))[0]
    tool_desc = TOOL_INFO.get(tool_id, ('', ''))[1]
    tone = (params or {}).get('tone', 'professional')
    strength = (params or {}).get('strength', 'strong')
    case = (params or {}).get('case', 'title')
    return f"""You are the writing engine inside TextForge Studio.

Task/tool: {tool_name}
Purpose: {tool_desc}
Tone: {tone}
Strength: {strength}
Case option when relevant: {case}

Rules:
1. Return only the final output. Do not explain unless the selected tool requires analysis.
2. Preserve the user's core meaning and factual claims.
3. Do not copy the original wording sentence-by-sentence. Change sentence structure, wording, transitions, and paragraph flow.
4. Improve grammar, punctuation, clarity, and readability.
5. Remove repetitive robotic phrasing and weak filler words.
6. For analysis tools, provide clear labels, practical scores, and short recommendations.
7. For rewriting/originality tools, create a polished, natural, human-readable draft. Encourage citations if the idea comes from another source.
8. Do not invent private facts, citations, statistics, or guarantees.

Input text:
{text}
"""


def call_gemini(api_key, model, prompt):
    if not api_key:
        raise ValueError('Gemini API key is missing. Add it in workspace settings.')
    model = model or 'gemini-1.5-flash'
    url = f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}'
    payload = {
        'contents': [{'parts': [{'text': prompt}]}],
        'generationConfig': {'temperature': 0.75, 'topP': 0.9, 'maxOutputTokens': 4096}
    }
    data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'})
    with urllib.request.urlopen(req, timeout=30) as res:
        body = json.loads(res.read().decode('utf-8'))
    try:
        return body['candidates'][0]['content']['parts'][0]['text'].strip()
    except Exception:
        raise RuntimeError('Gemini returned an unexpected response.')


def hybrid_process(tool_id, text, params, user_settings):
    algorithm_out = fallback_tool_output(tool_id, text, params)
    mode = (params or {}).get('engine_mode') or (user_settings['engine_mode'] if user_settings else 'algorithm')
    if mode != 'api':
        return algorithm_out, 'algorithm'
    api_key = (params or {}).get('gemini_api_key') or (user_settings['gemini_api_key'] if user_settings else '')
    model = (params or {}).get('gemini_model') or (user_settings['gemini_model'] if user_settings else 'gemini-1.5-flash')
    try:
        ai_out = call_gemini(api_key.strip(), model.strip(), make_gemini_prompt(tool_id, text, params))
        # Hybrid quality guard: if Gemini copies too much or returns too little, enforce local rewrite/polish.
        if not ai_out or (len(ai_out) < 0.25 * len(text) and len(text) > 200):
            return algorithm_out, 'algorithm-fallback'
        if tool_id in PREMIUM_TOOLS or 'rewrite' in tool_id or 'humanizer' in tool_id or 'cleaner' in tool_id:
            if jaccard_similarity(text, ai_out) > 75 or phrase_overlap(text, ai_out) > 70:
                ai_out = fallback_tool_output('deep_rewriter', ai_out, {'tone': params.get('tone','professional'), 'strength': 'maximum'})
                return grammar_polish(ai_out), 'hybrid-api-plus-algorithm'
        return grammar_polish(ai_out), 'gemini-api'
    except Exception as exc:
        return algorithm_out + f'\n\n[System note: API mode fell back to the built-in algorithm because the API request failed: {exc}]', 'algorithm-fallback'


@app.context_processor
def inject_globals():
    settings = site_settings()
    ads_enabled = settings.get('ads_enabled') == '1'
    user = current_user()
    return dict(
        app_name=APP_NAME,
        current_user=user,
        site_settings=settings,
        ads_enabled=ads_enabled,
        owner_name=OWNER_NAME,
        owner_email=OWNER_EMAIL,
        owner_phone=OWNER_PHONE,
        project_created_at=PROJECT_CREATED_AT,
        cat_slug=cat_slug,
        plan_status=plan_status(user, settings) if user else None,
        nav_tool_categories=TOOL_CATEGORIES,
        category_descriptions=CATEGORY_DESCRIPTIONS,
        feature_count=sum(len(ids) for ids in TOOL_CATEGORIES.values()),
    )


@app.route('/')
def index():
    return render_template('index.html', title=f'{APP_NAME} - Professional Writing Tools', categories=TOOL_CATEGORIES, tool_info=TOOL_INFO, guides=GUIDES)


@app.route('/about')
def about():
    return render_template('about.html', title='About TextForge Studio')


@app.route('/privacy-policy')
def privacy_policy():
    return render_template('privacy.html', title='Privacy Policy', legal_updated=LEGAL_UPDATED)


@app.route('/terms-and-conditions')
def terms():
    return render_template('terms.html', title='Terms and Conditions', legal_updated=LEGAL_UPDATED)


@app.route('/disclaimer')
def disclaimer():
    return render_template('disclaimer.html', title='Disclaimer', legal_updated=LEGAL_UPDATED)


@app.route('/guides')
def guides():
    return render_template('guides.html', title='Writing Guides', guides=GUIDES)


@app.route('/guides/<slug>')
def guide_detail(slug):
    guide = GUIDES.get(slug)
    if not guide:
        return redirect(url_for('guides'))
    return render_template('guide.html', title=guide['title'], guide=guide, slug=slug)


@app.route('/pricing')
def pricing():
    settings = site_settings()
    return render_template('pricing.html', title='Pricing', settings=settings)


@app.route('/signup', methods=['GET','POST'])
def signup():
    captcha_question = make_captcha() if request.method == 'GET' else session.get('captcha_question')
    if request.method == 'POST':
        if not check_captcha():
            flash('Captcha answer is incorrect. Please try again.', 'error')
            captcha_question = make_captcha()
            session['captcha_question'] = captcha_question
            return render_template('auth.html', mode='signup', captcha_question=captcha_question)
        name = request.form.get('name','').strip()
        email = request.form.get('email','').strip().lower()
        password = request.form.get('password','')
        if not name or not email or len(password) < 6:
            flash('Please enter your name, email, and a password with at least 6 characters.', 'error')
            captcha_question = make_captcha(); session['captcha_question'] = captcha_question
            return render_template('auth.html', mode='signup', captcha_question=captcha_question)
        try:
            now = datetime.utcnow().isoformat()
            with get_db() as db:
                cur = db.execute('INSERT INTO users(name,email,password_hash,created_at,role,plan,trial_started_at) VALUES(?,?,?,?,?,?,?)',
                    (name, email, generate_password_hash(password), now, 'user', 'free', now))
                user_id = cur.lastrowid
                db.execute('INSERT INTO settings(user_id) VALUES(?)', (user_id,))
            session['user_id'] = user_id
            flash('Account created. Your 1-day trial is active.', 'success')
            return redirect(url_for('dashboard'))
        except sqlite3.IntegrityError:
            flash('This email already has an account. Please login instead.', 'error')
    captcha_question = make_captcha(); session['captcha_question'] = captcha_question
    return render_template('auth.html', mode='signup', captcha_question=captcha_question)


@app.route('/login', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        if not check_captcha():
            flash('Captcha answer is incorrect. Please try again.', 'error')
            captcha_question = make_captcha(); session['captcha_question'] = captcha_question
            return render_template('auth.html', mode='login', captcha_question=captcha_question)
        email = request.form.get('email','').strip().lower()
        password = request.form.get('password','')
        with get_db() as db:
            user = db.execute('SELECT * FROM users WHERE email=?', (email,)).fetchone()
        if user and check_password_hash(user['password_hash'], password):
            session.clear()
            session['user_id'] = user['id']
            if user['role'] == 'admin':
                flash('Admin login successful. Admin Panel is now available on your dashboard.', 'success')
            return redirect(url_for('dashboard'))
        flash('Invalid email or password.', 'error')
    captcha_question = make_captcha(); session['captcha_question'] = captcha_question
    return render_template('auth.html', mode='login', captcha_question=captcha_question)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))


@app.route('/dashboard')
@login_required
def dashboard():
    return redirect(url_for('tools_home'))


@app.route('/dashboard/<category_slug>')
@login_required
def dashboard_category(category_slug):
    return redirect(url_for('tool_category', category_slug=category_slug))


@app.route('/tools')
@login_required
def tools_home():
    user = current_user()
    settings = site_settings()
    status = plan_status(user, settings)
    used = today_usage(user['id'])
    tools = build_tools_payload()
    return render_template('tool_categories.html',
                           title='Tool Categories',
                           user=user,
                           tools=tools,
                           user_plan=status,
                           usage_used=used,
                           usage_limit=status['limit'])


@app.route('/tools/<category_slug>')
@login_required
def tool_category(category_slug):
    user = current_user()
    selected_category = category_by_slug(category_slug)
    if not selected_category:
        return redirect(url_for('tools_home'))
    settings = site_settings()
    status = plan_status(user, settings)
    used = today_usage(user['id'])
    tools = build_tools_payload()
    return render_template('category_tools.html',
                           title=f'{selected_category} Tools',
                           user=user,
                           tools=tools,
                           selected_category=selected_category,
                           category_description=CATEGORY_DESCRIPTIONS.get(selected_category, 'Choose a tool and create polished content.'),
                           user_plan=status,
                           usage_used=used,
                           usage_limit=status['limit'])


@app.route('/tool/<tool_id>')
@login_required
def tool_page(tool_id):
    user = current_user()
    tool = find_tool(tool_id)
    if not tool:
        flash('Tool not found. Please choose another tool.', 'error')
        return redirect(url_for('tools_home'))
    settings = site_settings()
    status = plan_status(user, settings)
    used = today_usage(user['id'])
    with get_db() as db:
        user_settings = db.execute('SELECT * FROM settings WHERE user_id=?', (session['user_id'],)).fetchone()
        history = db.execute('SELECT * FROM history WHERE user_id=? ORDER BY id DESC LIMIT 12', (session['user_id'],)).fetchall()
    related = [find_tool(tid) for tid in TOOL_CATEGORIES.get(tool['category'], []) if tid != tool_id]
    related = [t for t in related if t][:6]
    return render_template('tool_page.html',
                           title=f'{tool["name"]} - TextForge Studio',
                           user=user,
                           tool=tool,
                           related=related,
                           history=history,
                           settings=user_settings,
                           tool_info=TOOL_INFO,
                           selected_category=tool['category'],
                           user_plan=status,
                           usage_used=used,
                           usage_limit=status['limit'])


@app.route('/admin', methods=['GET','POST'])
@login_required
@admin_required
def admin_panel():
    if request.method == 'POST':
        action = request.form.get('action', 'save_site')
        with get_db() as db:
            if action == 'save_site':
                allowed = ['ads_enabled','adsense_head_script','ad_header_html','ad_sidebar_html','ad_footer_html','ad_tool_html','ads_txt_publisher_id','site_domain','plan_price_usd','trial_days','free_daily_chars','trial_daily_chars','premium_daily_chars','premium_duration_days','enable_stripe','stripe_checkout_url','enable_lemon','lemon_checkout_url','enable_paddle','paddle_checkout_url','enable_paypal','paypal_details','enable_payoneer','payoneer_details','enable_bank','bank_details','enable_easypaisa','easypaisa_number','easypaisa_name','enable_sadapay','sadapay_number','sadapay_name','enable_mashreq','mashreq_details','enable_allied','allied_details','enable_api_mode','admin_notice']
                checkbox_keys = {'ads_enabled','enable_stripe','enable_lemon','enable_paddle','enable_paypal','enable_payoneer','enable_bank','enable_easypaisa','enable_sadapay','enable_mashreq','enable_allied','enable_api_mode'}
                for key in allowed:
                    value = request.form.get(key, '')
                    if key in checkbox_keys:
                        value = '1' if value == '1' else '0'
                    db.execute('UPDATE site_settings SET value=? WHERE key=?', (value, key))
                flash('Admin settings saved successfully.', 'success')
            elif action in ('approve_payment','reject_payment'):
                proof_id = request.form.get('proof_id')
                admin_note = request.form.get('admin_note','')
                proof = db.execute('SELECT * FROM payment_proofs WHERE id=?', (proof_id,)).fetchone()
                if proof:
                    status = 'approved' if action == 'approve_payment' else 'rejected'
                    db.execute('UPDATE payment_proofs SET status=?, reviewed_at=?, admin_note=? WHERE id=?', (status, datetime.utcnow().isoformat(), admin_note, proof_id))
                    if status == 'approved':
                        days = int_setting(site_settings(), 'premium_duration_days', 30)
                        until = (datetime.utcnow() + timedelta(days=days)).isoformat()
                        db.execute('UPDATE users SET plan=?, premium_until=? WHERE id=?', ('premium', until, proof['user_id']))
                    flash(f'Payment proof {status}.', 'success')
            elif action == 'update_user_plan':
                user_id = request.form.get('user_id')
                plan = request.form.get('plan','free')
                days = int(request.form.get('days','30') or '30')
                until = (datetime.utcnow() + timedelta(days=days)).isoformat() if plan == 'premium' else ''
                db.execute('UPDATE users SET plan=?, premium_until=? WHERE id=?', (plan, until, user_id))
                flash('User plan updated.', 'success')
        return redirect(url_for('admin_panel'))
    settings = site_settings()
    with get_db() as db:
        stats = {
            'users': db.execute('SELECT COUNT(*) AS c FROM users').fetchone()['c'],
            'premium': db.execute("SELECT COUNT(*) AS c FROM users WHERE plan='premium'").fetchone()['c'],
            'pending': db.execute("SELECT COUNT(*) AS c FROM payment_proofs WHERE status='pending'").fetchone()['c'],
            'history': db.execute('SELECT COUNT(*) AS c FROM history').fetchone()['c'],
        }
        users = db.execute('SELECT id,name,email,role,plan,created_at,premium_until FROM users ORDER BY id DESC LIMIT 100').fetchall()
        proofs = db.execute('SELECT p.*, u.name, u.email FROM payment_proofs p JOIN users u ON p.user_id=u.id ORDER BY p.id DESC LIMIT 100').fetchall()
    return render_template('admin.html', title='Admin Control Center', settings=settings, stats=stats, users=users, proofs=proofs)


@app.route('/subscribe', methods=['GET','POST'])
@login_required
def subscribe():
    settings = site_settings()
    user = current_user()
    if request.method == 'POST':
        method = request.form.get('method','manual')
        amount = request.form.get('amount', settings.get('plan_price_usd','5'))
        currency = request.form.get('currency','USD')
        transaction_ref = request.form.get('transaction_ref','')
        note = request.form.get('note','')
        file = request.files.get('screenshot')
        screenshot_path = ''
        if file and file.filename:
            os.makedirs(PAYMENT_UPLOAD_FOLDER, exist_ok=True)
            name = secure_filename(file.filename)
            ext = os.path.splitext(name)[1].lower()
            if ext not in {'.png','.jpg','.jpeg','.webp','.pdf'}:
                flash('Upload PNG, JPG, WEBP, or PDF payment proof only.', 'error')
                return redirect(url_for('subscribe'))
            filename = f'user{user["id"]}_{int(datetime.utcnow().timestamp())}_{name}'
            dest = os.path.join(PAYMENT_UPLOAD_FOLDER, filename)
            file.save(dest)
            screenshot_path = filename
        with get_db() as db:
            db.execute('INSERT INTO payment_proofs(user_id,method,amount,currency,transaction_ref,screenshot_path,note,status,created_at) VALUES(?,?,?,?,?,?,?,?,?)',
                (user['id'], method, amount, currency, transaction_ref, screenshot_path, note, 'pending', datetime.utcnow().isoformat()))
        flash('Payment proof submitted. Admin approval will upgrade your account automatically.', 'success')
        return redirect(url_for('subscribe'))
    with get_db() as db:
        proofs = db.execute('SELECT * FROM payment_proofs WHERE user_id=? ORDER BY id DESC LIMIT 10', (user['id'],)).fetchall()
    return render_template('subscribe.html', title='Upgrade Subscription', settings=settings, proofs=proofs, user_plan=plan_status(user, settings), usage_used=today_usage(user['id']))


@app.get('/payment-proof/<filename>')
@login_required
@admin_required
def payment_proof_file(filename):
    return send_file(os.path.join(PAYMENT_UPLOAD_FOLDER, secure_filename(filename)))


@app.post('/api/process')
@login_required
def api_process():
    data = request.get_json(force=True, silent=True) or {}
    tool_id = data.get('tool_id', 'humanizer')
    text = data.get('text', '') or ''
    params = data.get('params', {}) or {}
    if not text.strip():
        return jsonify({'ok': False, 'error': 'Please add text first.'}), 400
    user = current_user()
    site = site_settings()
    status = plan_status(user, site)
    used = today_usage(user['id'])
    requested_chars = len(text)
    if tool_id in PREMIUM_TOOLS and not (status['is_premium'] or status['is_trial'] or user['role'] == 'admin'):
        return jsonify({'ok': False, 'error': 'This is a premium tool. Upgrade to unlock it.', 'upgrade_required': True}), 402
    if used + requested_chars > status['limit']:
        return jsonify({'ok': False, 'error': f'Daily limit reached. Your plan allows {status["limit"]:,} characters per day. Upgrade for higher limits.', 'upgrade_required': True}), 402
    if params.get('engine_mode') == 'api' and not (status['is_premium'] or status['is_trial'] or user['role'] == 'admin'):
        return jsonify({'ok': False, 'error': 'Gemini API mode is available during trial or on premium plan.', 'upgrade_required': True}), 402
    user_settings = get_user_settings(user['id'])
    output, engine_used = hybrid_process(tool_id, text, params, user_settings)
    stats = metrics(text, output)
    add_usage(user['id'], requested_chars)
    if data.get('save', True):
        with get_db() as db:
            db.execute('INSERT INTO history(user_id,tool_id,input_text,output_text,stats_json,created_at,engine_used) VALUES(?,?,?,?,?,?,?)',
                (session['user_id'], tool_id, text[:10000], output[:15000], json.dumps(stats), datetime.utcnow().isoformat(), engine_used))
    return jsonify({'ok': True, 'output': output, 'stats': stats, 'tool': TOOL_INFO.get(tool_id, [tool_id])[0], 'engine_used': engine_used, 'usage': {'used': today_usage(user['id']), 'limit': status['limit']}})


@app.post('/api/upload')
@login_required
def api_upload():
    file = request.files.get('file')
    if not file:
        return jsonify({'ok': False, 'error': 'No file uploaded.'}), 400
    filename = secure_filename(file.filename or 'upload.txt')
    ext = os.path.splitext(filename)[1].lower()
    if ext not in UPLOAD_EXTENSIONS:
        return jsonify({'ok': False, 'error': 'Supported files: TXT, MD, DOCX, PDF.'}), 400
    raw = file.read()
    try:
        text = extract_text_from_upload(raw, ext)
        return jsonify({'ok': True, 'filename': filename, 'text': text})
    except Exception as exc:
        return jsonify({'ok': False, 'error': f'Could not read file: {exc}'}), 500


def extract_text_from_upload(raw, ext):
    if ext in ('.txt', '.md'):
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            return raw.decode('latin-1', errors='ignore')
    if ext == '.docx':
        from docx import Document
        doc = Document(io.BytesIO(raw))
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    if ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or '')
        return '\n\n'.join(pages).strip()
    return ''


@app.post('/api/export')
@login_required
def api_export():
    data = request.get_json(force=True, silent=True) or {}
    text = data.get('text','')
    fmt = data.get('format','txt').lower()
    title = data.get('title','TextForge Output')
    safe_title = ''.join(c for c in title if c.isalnum() or c in (' ','_','-')).strip()[:40] or 'textforge-output'
    if fmt == 'txt':
        bio = io.BytesIO(text.encode('utf-8'))
        return send_file(bio, as_attachment=True, download_name=f'{safe_title}.txt', mimetype='text/plain')
    if fmt == 'md':
        bio = io.BytesIO(text.encode('utf-8'))
        return send_file(bio, as_attachment=True, download_name=f'{safe_title}.md', mimetype='text/markdown')
    if fmt == 'docx':
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        for para in text.split('\n'):
            stripped = para.strip()
            if stripped.startswith('# '):
                doc.add_heading(stripped.strip('# ').strip(), level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped.strip('# ').strip(), level=2)
            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                doc.add_paragraph(para)
        bio = io.BytesIO(); doc.save(bio); bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=f'{safe_title}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    if fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=0.6*inch, leftMargin=0.6*inch, topMargin=0.6*inch, bottomMargin=0.6*inch)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles['Title']), Spacer(1, 12)]
        escaped = (text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'))
        for para in escaped.split('\n'):
            story.append(Paragraph(para if para.strip() else ' ', styles['BodyText']))
            story.append(Spacer(1, 6))
        doc.build(story); bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=f'{safe_title}.pdf', mimetype='application/pdf')
    return jsonify({'ok': False, 'error': 'Unsupported export format.'}), 400


@app.get('/api/history/<int:item_id>')
@login_required
def api_history_item(item_id):
    with get_db() as db:
        item = db.execute('SELECT * FROM history WHERE id=? AND user_id=?', (item_id, session['user_id'])).fetchone()
    if not item:
        return jsonify({'ok': False, 'error': 'Not found'}), 404
    return jsonify({'ok': True, 'item': dict(item)})


@app.post('/api/settings')
@login_required
def api_settings():
    data = request.get_json(force=True, silent=True) or {}
    tone = data.get('default_tone','professional')
    strength = data.get('default_strength','strong')
    brand = data.get('brand_name', APP_NAME)[:60]
    engine = data.get('engine_mode','algorithm')
    if engine not in ('algorithm','api'):
        engine = 'algorithm'
    api_key = data.get('gemini_api_key', None)
    model = data.get('gemini_model','gemini-1.5-flash')[:80]
    with get_db() as db:
        if api_key is None:
            db.execute('UPDATE settings SET default_tone=?, default_strength=?, brand_name=?, engine_mode=?, gemini_model=? WHERE user_id=?',
                       (tone, strength, brand, engine, model, session['user_id']))
        else:
            db.execute('UPDATE settings SET default_tone=?, default_strength=?, brand_name=?, engine_mode=?, gemini_api_key=?, gemini_model=? WHERE user_id=?',
                       (tone, strength, brand, engine, api_key.strip(), model, session['user_id']))
    return jsonify({'ok': True})


@app.post('/api/clear-history')
@login_required
def api_clear_history():
    with get_db() as db:
        db.execute('DELETE FROM history WHERE user_id=?', (session['user_id'],))
    return jsonify({'ok': True})


@app.route('/ads.txt')
def ads_txt():
    settings = site_settings()
    pub = settings.get('ads_txt_publisher_id','').strip()
    if pub and not pub.startswith('pub-'):
        pub = 'pub-' + pub
    if not pub:
        return Response('# Add your AdSense publisher ID from the admin panel.\n', mimetype='text/plain')
    return Response(f'google.com, {pub}, DIRECT, f08c47fec0942fa0\n', mimetype='text/plain')


@app.route('/robots.txt')
def robots_txt():
    domain = site_settings().get('site_domain','').rstrip('/')
    sitemap = f'Sitemap: {domain}/sitemap.xml\n' if domain else ''
    return Response(f'User-agent: *\nAllow: /\n{sitemap}', mimetype='text/plain')


@app.route('/sitemap.xml')
def sitemap_xml():
    domain = site_settings().get('site_domain','').rstrip('/') or request.url_root.rstrip('/')
    urls = ['', '/tools', '/about', '/pricing', '/privacy-policy', '/terms-and-conditions', '/disclaimer', '/guides'] + [f'/tools/{cat_slug(cat)}' for cat in TOOL_CATEGORIES] + [f'/guides/{slug}' for slug in GUIDES]
    items = ''.join(f'<url><loc>{domain}{u}</loc><changefreq>weekly</changefreq><priority>{"1.0" if u=="" else "0.8"}</priority></url>' for u in urls)
    return Response(f'<?xml version="1.0" encoding="UTF-8"?><urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">{items}</urlset>', mimetype='application/xml')


if __name__ == '__main__':
    app.run(debug=True)
